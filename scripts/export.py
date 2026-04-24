# -*- coding: utf-8 -*-
"""export.py — Multi-target exporter for questions.json.

Formats:
  anki    — CSV (or .apkg if genanki installed). Front=question, Back=answer+explanation, Tags=chapter type.
  quizlet — TSV, term TAB definition, blank line between cards.
  pdf     — Printable exam paper via reportlab; falls back to print-friendly HTML.
  docx    — Word document via python-docx.

Usage:
    python export.py --data questions.json --format {anki|quizlet|pdf|docx} --output <file>
"""
import argparse, csv, html, io, json, sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from common import get_version, utc_now_iso


# =====================================================================
# Helpers
# =====================================================================

def load_data(path: str) -> dict:
    p = Path(path)
    if not p.exists():
        sys.exit(f"ERROR: data file not found: {path}")
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        sys.exit(f"ERROR: invalid JSON: {e}")


def answer_text(q: dict) -> str:
    """Return human-readable answer string."""
    ans = q.get("answer", "")
    qtype = q.get("type", "single")
    options = q.get("options", {})
    if isinstance(ans, list):
        return "；".join(ans)
    if qtype in ("single", "multi", "tf") and options:
        parts = []
        for letter in ans.upper():
            if letter in options:
                parts.append(f"{letter}. {options[letter]}")
            else:
                parts.append(letter)
        return "；".join(parts) if parts else ans
    return str(ans)


def back_text(q: dict) -> str:
    """Anki back = answer text + optional explanation."""
    back = answer_text(q)
    expl = (q.get("explanation") or "").strip()
    if expl:
        back += f"\n\n解析：{expl}"
    return back


def tags_for(q: dict) -> str:
    chapter = (q.get("chapter") or "").replace(" ", "_")
    qtype = q.get("type", "single")
    return f"{chapter} {qtype}".strip()


# =====================================================================
# Anki
# =====================================================================

def export_anki_csv(questions: list[dict], output: str) -> None:
    out = Path(output)
    out.parent.mkdir(parents=True, exist_ok=True)
    with out.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Front", "Back", "Tags"])
        for q in questions:
            writer.writerow([
                q.get("question", ""),
                back_text(q),
                tags_for(q),
            ])
    print(f"Anki CSV written: {out}")
    print("Note: In Anki, import as 'Basic' note type (3 fields: Front, Back, Tags).")


def export_anki_apkg(questions: list[dict], output: str) -> None:
    try:
        import genanki  # type: ignore
    except ImportError:
        print("Warning: genanki not installed — falling back to CSV export.")
        csv_out = Path(output).with_suffix(".csv")
        export_anki_csv(questions, str(csv_out))
        return

    import random
    model_id = random.randrange(1 << 30, 1 << 31)
    deck_id = random.randrange(1 << 30, 1 << 31)

    model = genanki.Model(
        model_id,
        "quiz-gen Basic",
        fields=[{"name": "Front"}, {"name": "Back"}, {"name": "Tags"}],
        templates=[{
            "name": "Card 1",
            "qfmt": "{{Front}}",
            "afmt": "{{FrontSide}}<hr id=answer>{{Back}}",
        }],
    )
    deck = genanki.Deck(deck_id, "quiz-gen Export")
    for q in questions:
        note = genanki.Note(
            model=model,
            fields=[q.get("question", ""), back_text(q), tags_for(q)],
        )
        deck.add_note(note)

    out = Path(output)
    out.parent.mkdir(parents=True, exist_ok=True)
    genanki.Package(deck).write_to_file(str(out))
    print(f"Anki .apkg written: {out}")


def export_anki(questions: list[dict], output: str) -> None:
    if output.endswith(".apkg"):
        export_anki_apkg(questions, output)
    else:
        export_anki_csv(questions, output)


# =====================================================================
# Quizlet
# =====================================================================

def export_quizlet(questions: list[dict], output: str) -> None:
    out = Path(output)
    out.parent.mkdir(parents=True, exist_ok=True)
    lines = []
    for q in questions:
        term = q.get("question", "").replace("\t", " ").replace("\n", " ")
        definition = back_text(q).replace("\t", " ").replace("\n", " | ")
        lines.append(f"{term}\t{definition}")
        lines.append("")  # blank line between cards
    out.write_text("\n".join(lines), encoding="utf-8")
    print(f"Quizlet TSV written: {out}  ({len(questions)} cards)")


# =====================================================================
# PDF (reportlab or fallback HTML)
# =====================================================================

def export_pdf_reportlab(questions: list[dict], output: str) -> None:
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
    from reportlab.lib.units import cm  # type: ignore
    from reportlab.platypus import (  # type: ignore
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, PageBreak
    )
    from reportlab.lib.enums import TA_LEFT  # type: ignore
    from reportlab.lib import colors  # type: ignore
    import re

    out = Path(output)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(
        str(out),
        pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title2", parent=styles["Title"], fontSize=16, spaceAfter=12)
    q_style = ParagraphStyle("Q", parent=styles["Normal"], fontSize=11, spaceAfter=4, leading=16)
    opt_style = ParagraphStyle("Opt", parent=styles["Normal"], fontSize=10, spaceAfter=2,
                               leftIndent=20, leading=14)
    ans_style = ParagraphStyle("Ans", parent=styles["Normal"], fontSize=9, textColor=colors.grey)

    def safe(text: str) -> str:
        # Escape XML special chars for reportlab Paragraph
        return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = []
    story.append(Paragraph("题目预览 / Exam Paper", title_style))
    story.append(Spacer(1, 0.3 * cm))

    # Questions section
    for idx, q in enumerate(questions, 1):
        qtext = safe(q.get("question", ""))
        story.append(Paragraph(f"<b>{idx}.</b> [{q.get('type','?')}] {qtext}", q_style))
        for letter, text in sorted(q.get("options", {}).items()):
            story.append(Paragraph(f"{letter}. {safe(text)}", opt_style))
        story.append(Spacer(1, 0.4 * cm))

    # Answer key appendix
    story.append(PageBreak())
    story.append(Paragraph("答案 / Answer Key", title_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
    story.append(Spacer(1, 0.3 * cm))

    for idx, q in enumerate(questions, 1):
        ans_display = answer_text(q)
        expl = (q.get("explanation") or "").strip()
        text = f"<b>{idx}.</b> {safe(ans_display)}"
        if expl:
            text += f" — {safe(expl)}"
        story.append(Paragraph(text, ans_style))

    doc.build(story)
    print(f"PDF written: {out}")


def export_pdf_html_fallback(questions: list[dict], output: str) -> None:
    """Generate a print-friendly HTML when reportlab is unavailable."""
    out = Path(output)
    out.parent.mkdir(parents=True, exist_ok=True)

    q_rows = []
    ans_rows = []
    for idx, q in enumerate(questions, 1):
        qtext = html.escape(q.get("question", ""))
        opts_html = "".join(
            f'<div class="opt">{html.escape(letter)}. {html.escape(text)}</div>'
            for letter, text in sorted(q.get("options", {}).items())
        )
        q_rows.append(f"""
<div class="question">
  <p><strong>{idx}.</strong> [{html.escape(q.get('type','?'))}] {qtext}</p>
  {opts_html}
  <div class="answer-space"></div>
</div>""")

        ans_text = html.escape(answer_text(q))
        expl = html.escape((q.get("explanation") or "").strip())
        ans_part = f"<strong>{idx}.</strong> {ans_text}"
        if expl:
            ans_part += f" <span class='expl'>— {expl}</span>"
        ans_rows.append(f"<p>{ans_part}</p>")

    content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>Quiz Export</title>
  <style>
    body {{ font-family: serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
    .question {{ margin-bottom: 24px; page-break-inside: avoid; }}
    .opt {{ margin-left: 20px; }}
    .answer-space {{ height: 30px; border-bottom: 1px dashed #ccc; margin-top: 10px; }}
    .answer-key {{ page-break-before: always; }}
    .expl {{ color: #555; font-size: 0.9em; }}
    @media print {{ body {{ font-size: 12pt; }} }}
  </style>
</head>
<body>
<h1>题目 / Questions</h1>
{''.join(q_rows)}
<div class="answer-key">
  <h1>答案 / Answer Key</h1>
  {''.join(ans_rows)}
</div>
</body>
</html>"""
    out.write_text(content, encoding="utf-8")
    print(f"Print-friendly HTML written: {out}")
    print("Open in a browser and use File → Print → Save as PDF.")


def export_pdf(questions: list[dict], output: str) -> None:
    try:
        import reportlab  # type: ignore  # noqa: F401
        export_pdf_reportlab(questions, output)
    except ImportError:
        print("Warning: reportlab not installed — generating print-friendly HTML instead.")
        html_out = Path(output).with_suffix(".html")
        export_pdf_html_fallback(questions, str(html_out))


# =====================================================================
# DOCX
# =====================================================================

def export_docx(questions: list[dict], output: str) -> None:
    try:
        import docx as _docx  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except ImportError:
        sys.exit("ERROR: python-docx not installed. Run: pip install python-docx")

    document = _docx.Document()
    document.add_heading("题目 / Questions", level=0)

    for idx, q in enumerate(questions, 1):
        qtext = q.get("question", "")
        p = document.add_paragraph()
        run = p.add_run(f"{idx}. [{q.get('type', '?')}] {qtext}")
        run.bold = True

        for letter, text in sorted(q.get("options", {}).items()):
            op = document.add_paragraph(style="List Bullet")
            op.text = f"{letter}. {text}"

        document.add_paragraph()  # spacer

    # Answer key appendix
    document.add_page_break()
    document.add_heading("答案 / Answer Key", level=1)

    for idx, q in enumerate(questions, 1):
        ans = answer_text(q)
        expl = (q.get("explanation") or "").strip()
        line = f"{idx}. {ans}"
        if expl:
            line += f"  — {expl}"
        ap = document.add_paragraph(line)

    out = Path(output)
    out.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out))
    print(f"Word document written: {out}")


# =====================================================================
# CLI
# =====================================================================

FORMATS = {"anki", "quizlet", "pdf", "docx"}


def main():
    ap = argparse.ArgumentParser(description="Export questions.json to various formats.")
    ap.add_argument("--data", required=True, metavar="FILE",
                    help="Input questions.json")
    ap.add_argument("--format", required=True, choices=sorted(FORMATS), metavar="FORMAT",
                    help="Export format: anki | quizlet | pdf | docx")
    ap.add_argument("--output", required=True, metavar="FILE",
                    help="Output file path")
    args = ap.parse_args()

    data = load_data(args.data)
    questions = data.get("questions", [])
    if not questions:
        sys.exit("ERROR: no questions found in data file.")

    fmt = args.format
    out = args.output

    if fmt == "anki":
        export_anki(questions, out)
    elif fmt == "quizlet":
        export_quizlet(questions, out)
    elif fmt == "pdf":
        export_pdf(questions, out)
    elif fmt == "docx":
        export_docx(questions, out)
    else:
        sys.exit(f"ERROR: unknown format {fmt!r}")

    print(f"Done. {len(questions)} questions exported as {fmt}.")


if __name__ == "__main__":
    main()
